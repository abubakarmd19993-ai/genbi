import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def process_ocr(contents: bytes, filename: str, username: str) -> tuple:
    """Placeholder OCR - to be implemented later."""
    raw_text = "OCR processing will be available soon."
    doc = Document()
    doc.add_heading("GenBI OCR Document", 0)
    doc.add_paragraph(raw_text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read(), raw_text, raw_text