import io
import json
import fitz
import pandas as pd
from docx import Document as DocxDocument
from langchain_ollama import OllamaEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from datetime import datetime

embeddings = OllamaEmbeddings(model="nomic-embed-text")
CHROMA_DIR = "./chroma_db"

SUPPORTED_TYPES = {
    ".pdf": "PDF Document",
    ".docx": "Word Document",
    ".doc": "Word Document",
    ".txt": "Text File",
    ".csv": "CSV Dataset",
    ".xlsx": "Excel Dataset",
    ".json": "JSON File",
    ".md": "Markdown File",
}

def detect_file_type(filename: str) -> str:
    ext = "." + filename.lower().rsplit(".", 1)[-1]
    return SUPPORTED_TYPES.get(ext, "Unknown")

def extract_text(contents: bytes, filename: str) -> dict:
    ext = "." + filename.lower().rsplit(".", 1)[-1]

    if ext == ".pdf":
        pdf_doc = fitz.open(stream=contents, filetype="pdf")
        pages = []
        for i, page in enumerate(pdf_doc):
            text = page.get_text()
            if text.strip():
                pages.append({"page": i + 1, "text": text})
        pdf_doc.close()
        full_text = "\n".join([p["text"] for p in pages])
        return {"text": full_text, "pages": len(pages), "type": "pdf"}

    elif ext in [".docx", ".doc"]:
        doc = DocxDocument(io.BytesIO(contents))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)
        return {"text": full_text, "pages": 1, "paragraphs": len(paragraphs), "type": "docx"}

    elif ext == ".txt" or ext == ".md":
        text = contents.decode("utf-8", errors="ignore")
        return {"text": text, "pages": 1, "type": "txt"}

    elif ext == ".csv":
        df = pd.read_csv(io.BytesIO(contents))
        text = df.to_string(index=False)
        return {"text": text, "rows": len(df), "columns": len(df.columns), "cols": list(df.columns), "type": "csv"}

    elif ext == ".xlsx":
        xl = pd.ExcelFile(io.BytesIO(contents))
        all_text = []
        for sheet in xl.sheet_names:
            df = xl.parse(sheet)
            all_text.append(f"Sheet: {sheet}\n{df.to_string(index=False)}")
        text = "\n\n".join(all_text)
        return {"text": text, "sheets": len(xl.sheet_names), "sheet_names": xl.sheet_names, "type": "xlsx"}

    elif ext == ".json":
        data = json.loads(contents.decode("utf-8"))
        text = json.dumps(data, indent=2)
        return {"text": text, "type": "json"}

    else:
        raise ValueError(f"Unsupported file type: {ext}")

def embed_document(
    contents: bytes,
    filename: str,
    file_id: str,
    username: str,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
) -> dict:
    start_time = datetime.now()

    # Extract text
    doc_data = extract_text(contents, filename)
    text = doc_data["text"]

    if not text.strip():
        raise ValueError("No text found in document.")

    # Chunk text
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    chunks = splitter.split_text(text)

    # Create documents
    docs = [
        Document(
            page_content=chunk,
            metadata={
                "source": filename,
                "file_id": file_id,
                "username": username,
                "chunk_index": i,
                "doc_type": doc_data["type"],
                "embedded_at": datetime.now().isoformat(),
            }
        )
        for i, chunk in enumerate(chunks)
    ]

    # Store in ChromaDB
    collection_name = f"embed_{username[:10]}_{file_id[:15].replace('-', '_')}"
    vectorstore = Chroma(
        collection_name=collection_name,
        embedding_function=embeddings,
        persist_directory=CHROMA_DIR,
    )
    vectorstore.add_documents(docs)

    elapsed = (datetime.now() - start_time).seconds

    return {
        "file_id": file_id,
        "filename": filename,
        "file_type": detect_file_type(filename),
        "collection_name": collection_name,
        "chunks_created": len(chunks),
        "vectors_stored": len(chunks),
        "processing_time_seconds": elapsed,
        "embedded_at": datetime.now().strftime("%B %d, %Y %H:%M"),
        "metadata": {k: v for k, v in doc_data.items() if k != "text"},
        "status": "ready",
    }

def search_knowledge(
    query: str,
    collection_name: str,
    top_k: int = 4,
) -> dict:
    vectorstore = Chroma(
        collection_name=collection_name,
        embedding_function=embeddings,
        persist_directory=CHROMA_DIR,
    )
    results = vectorstore.similarity_search_with_relevance_scores(query, k=top_k)

    return {
        "query": query,
        "results": [
            {
                "content": doc.page_content,
                "source": doc.metadata.get("source", "Unknown"),
                "chunk_index": doc.metadata.get("chunk_index", 0),
                "relevance_score": round(score * 100, 1),
            }
            for doc, score in results
        ],
    }